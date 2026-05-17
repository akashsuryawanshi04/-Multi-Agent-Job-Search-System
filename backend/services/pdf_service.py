# ============================================================
# File: backend/services/pdf_service.py
# Purpose: Extract clean text from uploaded resume files.
#          Supports PDF (pdfplumber + PyMuPDF fallback) and
#          DOCX (python-docx). Handles edge cases like
#          multi-column layouts, scanned PDFs, and encoding
#          issues. Output feeds directly into Agent 1.
#
# Used by:
#   - backend/agents/resume_parser_agent.py → primary caller
#   - backend/api/routes/resume.py          → file validation
#
# Supported formats:
#   - .pdf  → pdfplumber (primary) + PyMuPDF (fallback)
#   - .docx → python-docx
#   - .doc  → python-docx with compatibility layer
# ============================================================

import io
import os
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple

from backend.config import settings
from backend.utils.logger import get_service_logger

logger = get_service_logger("PDFService")


# ============================================================
# Data Classes
# ============================================================

@dataclass
class ExtractionResult:
    """
    Result of a resume file extraction operation.

    Attributes:
        text:           Cleaned extracted text content.
        page_count:     Number of pages in the document.
        word_count:     Approximate word count of extracted text.
        method:         Extraction method used for logging/debugging.
        file_type:      Original file type ('pdf', 'docx', 'doc').
        warnings:       Non-fatal issues encountered during extraction.
        success:        True if extraction produced usable text.
        error:          Error message if extraction completely failed.
    """
    text: str
    page_count: int
    word_count: int
    method: str
    file_type: str
    warnings: List[str]
    success: bool
    error: Optional[str] = None

    @property
    def has_warnings(self) -> bool:
        """True if any non-fatal warnings were generated."""
        return len(self.warnings) > 0

    @property
    def is_likely_scanned(self) -> bool:
        """
        Heuristic to detect scanned/image-based PDFs.
        If word count is very low relative to page count,
        the PDF is likely a scanned image with no selectable text.
        """
        if self.page_count == 0:
            return False
        words_per_page = self.word_count / self.page_count
        return words_per_page < 50   # Threshold: fewer than 50 words/page

    def __repr__(self) -> str:
        return (
            f"<ExtractionResult "
            f"method={self.method} "
            f"pages={self.page_count} "
            f"words={self.word_count} "
            f"success={self.success}>"
        )


# ============================================================
# PDF Service Class
# ============================================================

class PDFService:
    """
    Service for extracting text from resume files.

    Extraction strategy per file type:

    PDF:
        1. Try pdfplumber — best for text-based PDFs,
           handles tables and multi-column layouts well.
        2. If pdfplumber yields < 100 words, fallback to
           PyMuPDF (fitz) which uses a different rendering
           engine and often succeeds where pdfplumber fails.
        3. If both fail, raise an informative error.

    DOCX:
        1. python-docx extracts paragraphs and tables.
        2. Preserves document structure better than PDF.

    After extraction:
        - Text is normalized (unicode, whitespace, encoding)
        - Boilerplate headers/footers are removed
        - Text is truncated if it exceeds Claude's context limit

    Usage:
        service = PDFService()
        result = await service.extract(file_path="/uploads/resume.pdf")

        if result.success:
            text = result.text
            # Send to Agent 1
        else:
            raise ValueError(result.error)
    """

    # Minimum word count to consider extraction successful
    MIN_WORD_COUNT = 20

    # Patterns to detect and remove common resume boilerplate
    # that adds noise without useful content
    _BOILERPLATE_PATTERNS = [
        r"page\s+\d+\s+of\s+\d+",           # "Page 1 of 3"
        r"^\s*\d+\s*$",                       # Lone page numbers
        r"curriculum\s+vitae\s*$",            # Standalone "Curriculum Vitae"
        r"confidential\s+resume",             # "Confidential Resume"
    ]

    def __init__(self) -> None:
        logger.info("PDFService initialized")

    # ----------------------------------------------------------
    # Primary Interface
    # ----------------------------------------------------------

    async def extract(
        self,
        file_path: str | Path,
        file_type: Optional[str] = None,
    ) -> ExtractionResult:
        """
        Extracts clean text from a resume file.

        Automatically detects file type from extension if not
        provided. Routes to the correct extraction method.

        Args:
            file_path: Absolute path to the uploaded resume file.
            file_type: Optional override ('pdf', 'docx', 'doc').
                       Auto-detected from extension if not given.

        Returns:
            ExtractionResult: Structured result with text and metadata.

        Raises:
            FileNotFoundError: If file_path does not exist.
            ValueError:        If file type is not supported.
        """
        path = Path(file_path)

        # Validate file exists
        if not path.exists():
            raise FileNotFoundError(
                f"Resume file not found at path: {file_path}"
            )

        # Detect file type
        detected_type = file_type or path.suffix.lstrip(".").lower()

        if detected_type not in ("pdf", "docx", "doc"):
            raise ValueError(
                f"Unsupported file type: '{detected_type}'. "
                f"Supported types: pdf, docx, doc"
            )

        logger.info(
            f"Starting extraction | "
            f"file={path.name} | "
            f"type={detected_type} | "
            f"size={path.stat().st_size:,} bytes"
        )

        # Route to correct extractor
        if detected_type == "pdf":
            result = await self._extract_pdf(path)
        else:
            result = await self._extract_docx(path, detected_type)

        # Post-process extracted text
        if result.success and result.text:
            result.text = self._clean_text(result.text)
            result.word_count = len(result.text.split())

            # Re-check minimum word count after cleaning
            if result.word_count < self.MIN_WORD_COUNT:
                result.success = False
                result.error = (
                    f"Extracted text is too short ({result.word_count} words). "
                    f"The file may be scanned, image-based, or password-protected. "
                    f"Please upload a text-based PDF or DOCX file."
                )

            # Warn if likely scanned
            if result.is_likely_scanned:
                result.warnings.append(
                    f"Low word density detected ({result.word_count} words "
                    f"across {result.page_count} pages). "
                    f"If this is a scanned resume, text extraction may be incomplete."
                )

            # Truncate to Claude's context limit
            result.text = self._truncate_for_context(result.text)

        logger.info(
            f"Extraction complete | "
            f"method={result.method} | "
            f"pages={result.page_count} | "
            f"words={result.word_count} | "
            f"success={result.success}"
        )

        if result.has_warnings:
            for warning in result.warnings:
                logger.warning(f"Extraction warning: {warning}")

        return result

    async def extract_from_bytes(
        self,
        file_bytes: bytes,
        filename: str,
    ) -> ExtractionResult:
        """
        Extracts text from file bytes (in-memory file).

        Used when the file comes from an HTTP upload before
        being saved to disk. Saves to a temporary location,
        extracts, then cleans up.

        Args:
            file_bytes: Raw bytes of the uploaded file.
            filename:   Original filename (used for type detection).

        Returns:
            ExtractionResult: Extraction result with text and metadata.
        """
        import tempfile

        suffix = Path(filename).suffix
        file_type = suffix.lstrip(".").lower()

        # Write to temp file for extraction
        with tempfile.NamedTemporaryFile(
            suffix=suffix,
            delete=False,
        ) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name

        try:
            result = await self.extract(
                file_path=tmp_path,
                file_type=file_type,
            )
        finally:
            # Always clean up temp file
            os.unlink(tmp_path)

        return result

    def validate_file(
        self,
        filename: str,
        file_size_bytes: int,
    ) -> Tuple[bool, Optional[str]]:
        """
        Validates a file before extraction.

        Checks file type and size against configured limits.
        Call this BEFORE saving the file to disk.

        Args:
            filename:        Original filename from the upload.
            file_size_bytes: Size of the file in bytes.

        Returns:
            Tuple[bool, Optional[str]]:
                (True, None) if valid.
                (False, error_message) if invalid.

        Usage:
            is_valid, error = service.validate_file(
                filename="resume.pdf",
                file_size_bytes=len(file_bytes),
            )
            if not is_valid:
                raise HTTPException(400, error)
        """
        # Check extension
        suffix = Path(filename).suffix.lstrip(".").lower()

        if suffix not in settings.allowed_file_types_list:
            return False, (
                f"File type '.{suffix}' is not allowed. "
                f"Please upload one of: "
                f"{', '.join(settings.allowed_file_types_list)}"
            )

        # Check size
        if file_size_bytes > settings.max_upload_size_bytes:
            max_mb = settings.max_upload_size_mb
            actual_mb = round(file_size_bytes / (1024 * 1024), 2)
            return False, (
                f"File size {actual_mb}MB exceeds the "
                f"{max_mb}MB limit. "
                f"Please compress or trim your resume."
            )

        # Check it's not empty
        if file_size_bytes == 0:
            return False, "File is empty. Please upload a valid resume."

        return True, None

    # ----------------------------------------------------------
    # PDF Extraction
    # ----------------------------------------------------------

    async def _extract_pdf(self, path: Path) -> ExtractionResult:
        """
        Extracts text from a PDF file.

        Tries pdfplumber first (better for structured PDFs),
        falls back to PyMuPDF if pdfplumber gets poor results.

        Args:
            path: Path to the PDF file.

        Returns:
            ExtractionResult with extracted text.
        """
        warnings = []

        # ── Strategy 1: pdfplumber ──────────────────────────────
        try:
            text, page_count = self._extract_with_pdfplumber(path)
            word_count = len(text.split())

            if word_count >= self.MIN_WORD_COUNT:
                logger.debug(
                    f"pdfplumber succeeded | "
                    f"pages={page_count} | words={word_count}"
                )
                return ExtractionResult(
                    text=text,
                    page_count=page_count,
                    word_count=word_count,
                    method="pdfplumber",
                    file_type="pdf",
                    warnings=warnings,
                    success=True,
                )
            else:
                warnings.append(
                    f"pdfplumber extracted only {word_count} words. "
                    f"Falling back to PyMuPDF."
                )
                logger.debug(
                    f"pdfplumber returned {word_count} words — "
                    f"trying PyMuPDF fallback"
                )

        except Exception as e:
            warnings.append(f"pdfplumber failed: {str(e)}")
            logger.warning(f"pdfplumber extraction failed: {e}")

        # ── Strategy 2: PyMuPDF (fitz) ─────────────────────────
        try:
            text, page_count = self._extract_with_pymupdf(path)
            word_count = len(text.split())

            if word_count >= self.MIN_WORD_COUNT:
                logger.debug(
                    f"PyMuPDF succeeded | "
                    f"pages={page_count} | words={word_count}"
                )
                return ExtractionResult(
                    text=text,
                    page_count=page_count,
                    word_count=word_count,
                    method="pymupdf",
                    file_type="pdf",
                    warnings=warnings,
                    success=True,
                )
            else:
                warnings.append(
                    f"PyMuPDF extracted only {word_count} words."
                )

        except Exception as e:
            warnings.append(f"PyMuPDF failed: {str(e)}")
            logger.warning(f"PyMuPDF extraction failed: {e}")

        # ── Both strategies failed ──────────────────────────────
        return ExtractionResult(
            text="",
            page_count=0,
            word_count=0,
            method="failed",
            file_type="pdf",
            warnings=warnings,
            success=False,
            error=(
                "Could not extract readable text from this PDF. "
                "The file may be scanned, image-based, "
                "or password-protected. "
                "Please upload a text-based PDF or DOCX file."
            ),
        )

    def _extract_with_pdfplumber(
        self,
        path: Path,
    ) -> Tuple[str, int]:
        """
        Extracts text using pdfplumber.

        pdfplumber is excellent at:
        - Standard text-based PDFs
        - PDFs with tables (extracts table content too)
        - Multi-column layouts (with bounding box strategy)

        Args:
            path: Path to the PDF file.

        Returns:
            Tuple[str, int]: (extracted_text, page_count)
        """
        import pdfplumber

        text_parts: List[str] = []
        page_count = 0

        with pdfplumber.open(str(path)) as pdf:
            page_count = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    # Primary extraction: standard text
                    page_text = page.extract_text(
                        x_tolerance=3,      # Horizontal character grouping
                        y_tolerance=3,      # Vertical line grouping
                        layout=True,        # Preserve spatial layout
                        x_density=7.25,     # Column detection density
                        y_density=13,       # Row detection density
                    )

                    if page_text:
                        text_parts.append(page_text)
                    else:
                        # Fallback: try extracting words individually
                        words = page.extract_words(
                            x_tolerance=5,
                            y_tolerance=5,
                            keep_blank_chars=False,
                        )
                        if words:
                            word_text = " ".join(
                                w["text"] for w in words
                            )
                            text_parts.append(word_text)

                    # Also extract table content if present
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            if table:
                                table_text = self._table_to_text(table)
                                if table_text:
                                    text_parts.append(table_text)

                except Exception as e:
                    logger.debug(
                        f"pdfplumber: error on page {page_num}: {e}"
                    )
                    continue

        return "\n\n".join(text_parts), page_count

    def _extract_with_pymupdf(
        self,
        path: Path,
    ) -> Tuple[str, int]:
        """
        Extracts text using PyMuPDF (fitz).

        PyMuPDF is excellent at:
        - PDFs with embedded fonts
        - Complex layouts that confuse pdfplumber
        - Preserving reading order

        Args:
            path: Path to the PDF file.

        Returns:
            Tuple[str, int]: (extracted_text, page_count)
        """
        import fitz   # PyMuPDF

        text_parts: List[str] = []
        page_count = 0

        doc = fitz.open(str(path))

        try:
            page_count = len(doc)

            for page_num in range(page_count):
                try:
                    page = doc[page_num]

                    # Extract text preserving reading order
                    # "blocks" mode groups text by visual blocks
                    blocks = page.get_text("blocks", sort=True)

                    page_texts = []
                    for block in blocks:
                        # block = (x0, y0, x1, y1, text, block_no, block_type)
                        if len(block) >= 5 and block[6] == 0:  # type 0 = text
                            block_text = block[4].strip()
                            if block_text:
                                page_texts.append(block_text)

                    if page_texts:
                        text_parts.append("\n".join(page_texts))

                except Exception as e:
                    logger.debug(
                        f"PyMuPDF: error on page {page_num + 1}: {e}"
                    )
                    continue

        finally:
            doc.close()

        return "\n\n".join(text_parts), page_count

    # ----------------------------------------------------------
    # DOCX Extraction
    # ----------------------------------------------------------

    async def _extract_docx(
        self,
        path: Path,
        file_type: str,
    ) -> ExtractionResult:
        """
        Extracts text from a DOCX (or DOC) file.

        Extracts:
        - All paragraphs (body text, headings, bullets)
        - Table cell content
        - Text boxes (via XML fallback)
        - Header and footer content

        Args:
            path:      Path to the DOCX file.
            file_type: 'docx' or 'doc'.

        Returns:
            ExtractionResult with extracted text.
        """
        warnings = []

        try:
            text, page_count = self._extract_with_python_docx(
                path, file_type
            )
            word_count = len(text.split())

            if word_count < self.MIN_WORD_COUNT:
                return ExtractionResult(
                    text=text,
                    page_count=page_count,
                    word_count=word_count,
                    method="python-docx",
                    file_type=file_type,
                    warnings=warnings,
                    success=False,
                    error=(
                        f"DOCX extraction yielded only {word_count} words. "
                        f"The file may be corrupted or contain only images."
                    ),
                )

            return ExtractionResult(
                text=text,
                page_count=page_count,
                word_count=word_count,
                method="python-docx",
                file_type=file_type,
                warnings=warnings,
                success=True,
            )

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ExtractionResult(
                text="",
                page_count=0,
                word_count=0,
                method="failed",
                file_type=file_type,
                warnings=warnings,
                success=False,
                error=(
                    f"Failed to extract text from DOCX file: {str(e)}. "
                    f"The file may be corrupted, password-protected, "
                    f"or saved in an incompatible format."
                ),
            )

    def _extract_with_python_docx(
        self,
        path: Path,
        file_type: str,
    ) -> Tuple[str, int]:
        """
        Extracts text from DOCX using python-docx.

        Handles:
        - Regular paragraphs and headings
        - Bullet points and numbered lists
        - Tables (cell by cell)
        - Headers and footers
        - Text boxes via XML namespace

        Args:
            path:      Path to the DOCX file.
            file_type: 'docx' or 'doc'.

        Returns:
            Tuple[str, int]: (extracted_text, estimated_page_count)
        """
        import docx
        from docx.oxml.ns import qn

        doc = docx.Document(str(path))
        text_parts: List[str] = []

        # ── Extract paragraphs ──────────────────────────────────
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                text_parts.append(para_text)

        # ── Extract tables ──────────────────────────────────────
        for table in doc.tables:
            table_rows: List[str] = []
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                text_parts.append("\n".join(table_rows))

        # ── Extract headers and footers ─────────────────────────
        for section in doc.sections:
            try:
                # Header
                header = section.header
                if header:
                    for para in header.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text.strip())

                # Footer
                footer = section.footer
                if footer:
                    for para in footer.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text.strip())
            except Exception:
                pass   # Headers/footers optional — skip if unavailable

        # ── Extract text boxes via XML ──────────────────────────
        # Text boxes in DOCX are stored in XML drawing elements
        try:
            body = doc.element.body
            for textbox in body.iter(qn("wps:txbx")):
                for para in textbox.iter(qn("w:p")):
                    texts = [
                        node.text
                        for node in para.iter(qn("w:t"))
                        if node.text
                    ]
                    combined = "".join(texts).strip()
                    if combined:
                        text_parts.append(combined)
        except Exception:
            pass   # Text boxes optional — skip if namespace unavailable

        full_text = "\n\n".join(text_parts)

        # Estimate page count from word count
        # ~250 words per page is a reasonable approximation
        word_count = len(full_text.split())
        estimated_pages = max(1, word_count // 250)

        return full_text, estimated_pages

    # ----------------------------------------------------------
    # Text Cleaning
    # ----------------------------------------------------------

    def _clean_text(self, text: str) -> str:
        """
        Cleans and normalizes extracted resume text.

        Operations (in order):
            1. Unicode normalization (handles special chars)
            2. Remove null bytes and control characters
            3. Normalize unicode quotes and dashes
            4. Fix common encoding artifacts (â€™ → ')
            5. Remove boilerplate patterns (page numbers etc.)
            6. Normalize whitespace (collapse multiple spaces/newlines)
            7. Strip leading/trailing whitespace

        Args:
            text: Raw extracted text from PDF or DOCX.

        Returns:
            str: Cleaned, normalized text ready for Claude.
        """
        if not text:
            return ""

        # Step 1: Unicode normalization (NFC = composed form)
        text = unicodedata.normalize("NFC", text)

        # Step 2: Remove null bytes and non-printable control chars
        # Keep: tab (\t), newline (\n), carriage return (\r)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

        # Step 3: Normalize unicode quotes and dashes to ASCII
        replacements = {
            "\u2018": "'",   # Left single quotation mark
            "\u2019": "'",   # Right single quotation mark
            "\u201c": '"',   # Left double quotation mark
            "\u201d": '"',   # Right double quotation mark
            "\u2013": "-",   # En dash
            "\u2014": "-",   # Em dash
            "\u2022": "-",   # Bullet point
            "\u2023": "-",   # Triangular bullet
            "\u25cf": "-",   # Black circle bullet
            "\u25aa": "-",   # Black small square bullet
            "\u2212": "-",   # Minus sign
            "\u00a0": " ",   # Non-breaking space
            "\u200b": "",    # Zero-width space
            "\ufeff": "",    # Byte order mark
        }
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)

        # Step 4: Fix common UTF-8 encoding artifacts
        # These appear when text is double-encoded
        encoding_fixes = {
            "â€™": "'",
            "â€œ": '"',
            "â€": '"',
            "â€"": "-",
            "â€"": "-",
            "Ã©": "é",
            "Ã¨": "è",
            "Ã ": "à",
        }
        for artifact, fix in encoding_fixes.items():
            text = text.replace(artifact, fix)

        # Step 5: Remove boilerplate patterns line by line
        lines = text.split("\n")
        cleaned_lines = []
        for line in lines:
            stripped = line.strip()
            is_boilerplate = any(
                re.search(pattern, stripped, re.IGNORECASE)
                for pattern in self._BOILERPLATE_PATTERNS
            )
            if not is_boilerplate:
                cleaned_lines.append(line)
        text = "\n".join(cleaned_lines)

        # Step 6: Normalize whitespace
        # Collapse multiple spaces to single space
        text = re.sub(r"[ \t]+", " ", text)

        # Collapse 3+ consecutive newlines to 2 newlines max
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Remove trailing whitespace from each line
        text = "\n".join(line.rstrip() for line in text.split("\n"))

        # Step 7: Final strip
        return text.strip()

    def _table_to_text(
        self,
        table: List[List[Optional[str]]],
    ) -> str:
        """
        Converts a pdfplumber table to readable text.

        Joins non-empty cells with " | " separator
        and rows with newlines.

        Args:
            table: 2D list of cell values from pdfplumber.

        Returns:
            str: Human-readable table text.
        """
        rows = []
        for row in table:
            if row:
                cells = [
                    str(cell).strip()
                    for cell in row
                    if cell and str(cell).strip()
                ]
                if cells:
                    rows.append(" | ".join(cells))
        return "\n".join(rows)

    def _truncate_for_context(
        self,
        text: str,
        max_chars: int = 80_000,
    ) -> str:
        """
        Truncates text to fit within Claude's context window.

        80,000 characters ≈ 20,000 tokens, which leaves plenty
        of room for the system prompt and response.

        For resumes, this limit is almost never hit (resumes are
        typically 500-2000 words = 3000-12000 chars). This is a
        safety net for unusually long documents.

        Args:
            text:      Text to potentially truncate.
            max_chars: Maximum characters to allow.

        Returns:
            str: Possibly truncated text.
        """
        if len(text) <= max_chars:
            return text

        logger.warning(
            f"Resume text truncated: {len(text)} → {max_chars} chars"
        )

        truncated = text[:max_chars]

        # Try to truncate at a word boundary
        last_space = truncated.rfind(" ", max_chars - 200)
        if last_space > max_chars - 500:
            truncated = truncated[:last_space]

        return (
            truncated
            + "\n\n[Resume truncated — content beyond this point was omitted]"
        )

    # ----------------------------------------------------------
    # Utility Methods
    # ----------------------------------------------------------

    def get_file_metadata(self, file_path: str | Path) -> dict:
        """
        Returns basic metadata about a resume file without
        performing full text extraction.

        Used for quick validation and display in the UI.

        Args:
            file_path: Path to the file.

        Returns:
            dict: File metadata including size, type, name.
        """
        path = Path(file_path)

        if not path.exists():
            return {"error": f"File not found: {file_path}"}

        stat = path.stat()
        return {
            "filename": path.name,
            "file_type": path.suffix.lstrip(".").lower(),
            "file_size_bytes": stat.st_size,
            "file_size_mb": round(stat.st_size / (1024 * 1024), 3),
            "created_at": stat.st_ctime,
            "modified_at": stat.st_mtime,
        }

    def is_supported_file(self, filename: str) -> bool:
        """
        Returns True if the file extension is supported.

        Args:
            filename: File name with extension.

        Returns:
            bool: True if file type is supported.
        """
        suffix = Path(filename).suffix.lstrip(".").lower()
        return suffix in ("pdf", "docx", "doc")


# ============================================================
# Module-Level Singleton
# ============================================================

pdf_service: PDFService = PDFService()